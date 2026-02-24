from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table_from_lines(doc, lines):
    if not lines:
        return
    # Parse markdown table
    rows = []
    for line in lines:
        line = line.strip()
        if not line or set(line.replace('|','').replace('-','').replace(' ','')) == set():
            continue  # skip separator rows
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= max_cols:
                break
            cell = table.rows[i].cells[j]
            # strip markdown bold
            cell_text = cell_text.replace('**', '').replace('*', '')
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                if i == 0:
                    for run in para.runs:
                        run.font.bold = True
    doc.add_paragraph()

def process_inline(para, text):
    """Add a paragraph with inline bold/italic/code formatting."""
    # Split on bold (**...**), italic (*...*), and inline code (`...`)
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = para.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            run = para.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        raw = lines[i].rstrip('\n').rstrip('\r')
        stripped = raw.strip()

        # Code block toggle
        if stripped.startswith('```'):
            if in_code_block:
                # End code block — write it
                if code_lines:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(1)
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after = Pt(4)
                    run = para.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                # flush any pending table
                if in_table:
                    add_table_from_lines(doc, table_lines)
                    table_lines = []
                    in_table = False
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # Table detection
        if stripped.startswith('|'):
            table_lines.append(stripped)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                add_table_from_lines(doc, table_lines)
                table_lines = []
                in_table = False

        # Horizontal rule
        if stripped in ('---', '___', '***'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
            i += 1
            continue

        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:].strip()
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith('### ') and not stripped.startswith('#### '):
            text = stripped[4:].strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x10, 0x36, 0x78)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.italic = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            text = stripped[2:].replace('> ', '').strip()
            text = text.replace('**', '').replace('*', '').replace('`', '')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* ') and len(stripped) > 2:
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            process_inline(p, text)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            process_inline(p, text)
            i += 1
            continue

        # Bold/author lines (lines starting with **)
        if stripped.startswith('**') and stripped.endswith('**') and '\n' not in stripped:
            text = stripped[2:-2]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            i += 1
            continue

        # Empty line
        if not stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(0.5)
        process_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')

if __name__ == '__main__':
    convert_md_to_docx(
        r'C:\Users\aloob\Downloads\Research Backup\Draft\SAKE_research_paper.md',
        r'C:\Users\aloob\Downloads\Research Backup\Draft\SAKE_research_paper.docx'
    )
