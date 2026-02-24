"""
build_clean_docx.py  —  Convert SAKE_research_paper_clean.md to DOCX with figures.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

MD  = r'C:\Users\aloob\Downloads\Research Backup\Draft\SAKE_research_paper_clean.md'
OUT = r'C:\Users\aloob\Downloads\Research Backup\Draft\SAKE_research_paper_clean.docx'
IMG = r'C:\Users\aloob\Downloads\Research Backup\simulation\results'

FIGURES = [
    (
        'sim_latency.png',
        'Fig. 1. Computational latency comparison: per-packet base paper QC-LDPC cost vs. SAKE Tier 2 AEAD. Break-even at N = 4; total savings reach 24.7x at N = 100.',
        'Key Result: 97.3',
        Inches(5.8),
    ),
    (
        'sim_bandwidth_bar.png',
        'Fig. 2. Per-packet bandwidth overhead (bar chart): CT0 = 408 bits vs. Tier 2 Nonce + TAG = 224 bits. Absolute saving: 184 bits/packet.',
        'Cumulative Bandwidth Comparison',
        Inches(5.0),
    ),
    (
        'sim_bandwidth_cumulative.png',
        'Fig. 3. Cumulative bandwidth overhead comparison over N packets. The per-packet gap of 184 bits widens linearly, converging to 45.1% overhead reduction asymptotically.',
        'Key Result: The 184 bits',
        Inches(5.0),
    ),
    (
        'sim_energy.png',
        'Fig. 4. Per-packet clock cycle comparison: SAKE Tier 2 (0.074x10^6) vs. Lizard, RLizard, LEDAkem, and base paper code-based HE. SAKE achieves the lowest count at all estimates.',
        'Key Result: Lowest clock',
        Inches(5.8),
    ),
]


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def process_inline(para, text):
    """Render **bold**, *italic*, and `code` inline runs."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    for part in pattern.split(text):
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            r = para.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
            continue
        else:
            r = para.add_run(part)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)


def add_table_from_lines(doc, lines):
    """Render a markdown table into a Word table with grid borders."""
    parsed = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        # Skip pure-separator rows (--- cells)
        if all(re.fullmatch(r'-+', c.replace(' ', '')) for c in cells if c):
            continue
        parsed.append(cells)
    if not parsed:
        return
    max_cols = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=max_cols)
    tbl.style = 'Table Grid'
    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row[:max_cols]):
            cell = tbl.rows[i].cells[j]
            clean = cell_text.replace('**', '').replace('*', '').replace('`', '')
            cell.text = clean
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    if i == 0:
                        run.bold = True
    doc.add_paragraph()


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_figure(doc, img_path, caption_text, width):
    """Add a centered image and italic IEEE-style caption; return both paragraphs."""
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(10)
    ip.paragraph_format.space_after  = Pt(4)
    ip.add_run().add_picture(img_path, width=width)

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after  = Pt(12)
    cr = cp.add_run(caption_text)
    cr.italic = True
    cr.font.name  = 'Times New Roman'
    cr.font.size  = Pt(10)
    cr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return ip, cp


def move_after(doc, anchor, new_paras):
    """Move new_paras so they appear immediately after the paragraph containing anchor."""
    body  = doc.element.body
    paras = body.findall(qn('w:p'))
    target = None
    for p_el in paras:
        txt = ''.join(r.text for r in p_el.findall('.//' + qn('w:t')) if r.text)
        if anchor.lower() in txt.lower():
            target = p_el
            break
    if target is None:
        return False
    ref = target
    for np in new_paras:
        ref.addnext(np._p)
        ref = np._p
    return True


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → DOCX conversion
# ─────────────────────────────────────────────────────────────────────────────

def convert(md_path, out_path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code    = False
    code_lines = []
    tbl_lines  = []
    in_tbl     = False

    while i < len(lines):
        raw = lines[i].rstrip()
        s   = raw.strip()

        # ── code block ──
        if s.startswith('```'):
            if in_code:
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent   = Cm(1)
                    p.paragraph_format.space_before  = Pt(4)
                    p.paragraph_format.space_after   = Pt(4)
                    r = p.add_run('\n'.join(code_lines))
                    r.font.name = 'Courier New'
                    r.font.size = Pt(9)
                code_lines = []
                in_code    = False
            else:
                if in_tbl:
                    add_table_from_lines(doc, tbl_lines)
                    tbl_lines = []
                    in_tbl    = False
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── table ──
        if s.startswith('|'):
            tbl_lines.append(s)
            in_tbl = True
            i += 1
            continue
        else:
            if in_tbl:
                add_table_from_lines(doc, tbl_lines)
                tbl_lines = []
                in_tbl    = False

        # ── horizontal rule ──
        if s in ('---', '___', '***'):
            add_hr(doc)
            i += 1
            continue

        # ── headings ──
        if s.startswith('# ') and not s.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(s[2:])
            r.bold = True
            r.font.name  = 'Times New Roman'
            r.font.size  = Pt(15)
            r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
            i += 1
            continue

        if s.startswith('## ') and not s.startswith('### '):
            p = doc.add_heading(s[3:], level=1)
            for r in p.runs:
                r.font.name  = 'Times New Roman'
                r.font.size  = Pt(13)
                r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        if s.startswith('### ') and not s.startswith('#### '):
            p = doc.add_heading(s[4:], level=2)
            for r in p.runs:
                r.font.name  = 'Times New Roman'
                r.font.size  = Pt(12)
                r.font.color.rgb = RGBColor(0x10, 0x36, 0x78)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(3)
            i += 1
            continue

        if s.startswith('#### '):
            p = doc.add_heading(s[5:], level=3)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
                r.italic    = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            i += 1
            continue

        # ── blockquote ──
        if s.startswith('> '):
            txt = s[2:].replace('**', '').replace('*', '').replace('`', '')
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(txt)
            r.italic = True
            r.font.name  = 'Times New Roman'
            r.font.size  = Pt(11)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # ── bullet list ──
        if s.startswith('- ') or (s.startswith('* ') and len(s) > 2):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent   = Cm(0.8)
            p.paragraph_format.space_before  = Pt(1)
            p.paragraph_format.space_after   = Pt(1)
            process_inline(p, s[2:])
            i += 1
            continue

        # ── numbered list ──
        if re.match(r'^\d+\. ', s):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent   = Cm(0.8)
            p.paragraph_format.space_before  = Pt(1)
            p.paragraph_format.space_after   = Pt(1)
            process_inline(p, re.sub(r'^\d+\. ', '', s))
            i += 1
            continue

        # ── blank line ──
        if not s:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── normal paragraph ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before       = Pt(2)
        p.paragraph_format.space_after        = Pt(4)
        p.paragraph_format.first_line_indent  = Cm(0.5)
        process_inline(p, s)
        i += 1

    # ── embed figures ──
    for fname, caption, anchor, width in FIGURES:
        img_path = os.path.join(IMG, fname)
        if not os.path.exists(img_path):
            print('MISSING:', img_path)
            continue
        ip, cp = add_figure(doc, img_path, caption, width)
        placed = move_after(doc, anchor, [ip, cp])
        status = 'placed after anchor' if placed else 'appended at end'
        print(fname + ': ' + status)

    doc.save(out_path)
    print('Saved:', out_path)


if __name__ == '__main__':
    convert(MD, OUT)
