"""
insert_figures.py
Inserts the 4 simulation graphs into SAKE_research_paper.docx
at the correct positions inside Section VI (Performance Evaluation).
Each figure is centered, full-width, with an IEEE-style caption below it.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

IMG_DIR = r'C:\Users\aloob\Downloads\Research Backup\simulation\results'
DOCX_IN  = r'C:\Users\aloob\Downloads\Research Backup\Draft\SAKE_research_paper.docx'
DOCX_OUT = r'C:\Users\aloob\Downloads\Research Backup\Draft\SAKE_research_paper_final.docx'

# Figure definitions: (filename, caption, anchor_text)
# anchor_text = a unique string that appears in the paragraph AFTER which the figure is inserted
FIGURES = [
    {
        'file': 'sim_latency.png',
        'caption': 'Fig. 1. Computation latency comparison: Base paper (QC-LDPC per packet) vs. proposed SAKE '
                   '(Epoch initiation + Tier 2 AEAD). Break-even at N = 4; 24.7× savings at N = 100. '
                   '97.3%–99.0% per-packet Tier 2 reduction.',
        'anchor': 'Key Result: 97.3',   # text near the end of §VI-A
        'width': Inches(5.8),
    },
    {
        'file': 'sim_bandwidth_bar.png',
        'caption': 'Fig. 2. Per-packet bandwidth overhead comparison (bar chart): Base paper CT₀ (408 bits) '
                   'vs. proposed Tier 2 AEAD overhead (96-bit Nonce + 128-bit TAG = 224 bits). '
                   'Absolute saving: 184 bits/packet.',
        'anchor': 'Cumulative Bandwidth Comparison',  # heading of the cumulative table
        'width': Inches(5.0),
    },
    {
        'file': 'sim_bandwidth_cumulative.png',
        'caption': 'Fig. 3. Cumulative bandwidth overhead comparison: Both curves start from an identical '
                   '27,592-bit epoch baseline (LR-IoTA + pk_HE). The proposed curve grows 45.1% more slowly '
                   'per packet, asymptotically saving 45.1% of per-packet overhead at large N.',
        'anchor': 'Key Result: 184 bits',   # end of §VI-B key result
        'width': Inches(5.0),
    },
    {
        'file': 'sim_energy.png',
        'caption': 'Fig. 4. Clock cycle comparison per data packet: Proposed SAKE Tier 2 (0.074×10⁶ cycles, '
                   'benchmark-estimated) vs. Original Lizard, RLizard, LEDAkem, and Base Paper Code-based HE. '
                   'SAKE achieves the lowest cycle count at both benchmark (33×) and worst-case (24×) estimates.',
        'anchor': 'Key Result: Lowest clock',   # end of §VI-C
        'width': Inches(5.8),
    },
]


def add_figure_paragraph(doc, img_path, caption_text, width):
    """Insert one centered image + caption as two paragraphs and return the first para."""
    # --- Image paragraph ---
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(8)
    img_para.paragraph_format.space_after  = Pt(4)
    run = img_para.add_run()
    run.add_picture(img_path, width=width)

    # --- Caption paragraph ---
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_before = Pt(2)
    cap_para.paragraph_format.space_after  = Pt(10)
    run = cap_para.add_run(caption_text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    return img_para, cap_para


def move_paragraphs_after(doc, anchor_text, new_paras):
    """
    Find the first paragraph whose text contains `anchor_text`,
    then insert `new_paras` immediately after it in the document body.
    Returns True if anchor was found.
    """
    body = doc.element.body
    paras = body.findall(qn('w:p'))
    target_p = None

    # Search all paragraphs for the anchor text
    for p_el in paras:
        full_text = ''.join(
            r.text for r in p_el.findall('.//' + qn('w:t')) if r.text
        )
        if anchor_text.lower() in full_text.lower():
            target_p = p_el
            break

    if target_p is None:
        print(f'  [WARN] Anchor not found: "{anchor_text}" — figure appended at end.')
        return False

    # Insert new paragraphs' XML elements right after target_p
    ref = target_p
    for new_p in new_paras:
        new_p_el = new_p._p
        ref.addnext(new_p_el)
        ref = new_p_el

    return True


def main():
    doc = Document(DOCX_IN)

    for fig in FIGURES:
        img_path = os.path.join(IMG_DIR, fig['file'])
        print(f'Inserting {fig["file"]} ...')

        if not os.path.exists(img_path):
            print(f'  [ERROR] Image not found: {img_path}')
            continue

        # Add the figure paragraphs at the END first (python-docx limitation —
        # we add then move them via XML manipulation)
        img_para, cap_para = add_figure_paragraph(doc, img_path, fig['caption'], fig['width'])

        # Now move those two paragraphs to the correct position
        found = move_paragraphs_after(doc, fig['anchor'], [img_para, cap_para])
        if found:
            print(f'  -> Placed after: "{fig["anchor"]}"')
        else:
            print(f'  -> Placed at end of document.')

    doc.save(DOCX_OUT)
    print(f'\nSaved: {DOCX_OUT}')


if __name__ == '__main__':
    main()
